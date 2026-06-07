"""将Markdown实训报告转为Word (.docx) 格式"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

INPUT = "人工智能专业综合实训报告.md"
OUTPUT = "人工智能专业综合实训报告.docx"


def set_cell_font(cell, text, bold=False, size=9, color=None, font_name="微软雅黑"):
    """设置单元格字体"""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_styled_paragraph(doc, text, style="Normal", bold=False, size=11, font_name="微软雅黑"):
    """添加带格式的段落"""
    p = doc.add_paragraph(style=style)
    # 处理粗体标记 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(size)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def parse_markdown_table(lines):
    """解析Markdown表格，返回表头和数据行"""
    if len(lines) < 2:
        return None, None
    # 过滤分隔行
    content_lines = [l for l in lines if not re.match(r'^\|[\s\-:|]+\|$', l)]
    if len(content_lines) < 1:
        return None, None

    def split_row(line):
        cells = line.strip().strip('|').split('|')
        return [c.strip() for c in cells]

    header = split_row(content_lines[0])
    rows = [split_row(l) for l in content_lines[1:]]
    return header, rows


def convert_md_to_docx(input_path, output_path):
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    table_buffer = []
    in_table = False
    in_code = False
    list_buffer = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 代码块
        if line.startswith('```'):
            if in_code:
                in_code = False
                i += 1
                continue
            else:
                in_code = True
                i += 1
                continue

        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.name = "Consolas"
            i += 1
            continue

        # 表格处理
        if line.startswith('|') and line.endswith('|'):
            table_buffer.append(line)
            in_table = True
            i += 1
            # 检查下一行是否还在表格中
            if i < len(lines) and lines[i].startswith('|'):
                continue

        if in_table:
            header, rows = parse_markdown_table(table_buffer)
            if header and rows:
                table = doc.add_table(rows=len(rows) + 1, cols=len(header), style='Light Grid Accent 1')
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 表头
                for j, h in enumerate(header):
                    set_cell_font(table.rows[0].cells[j], h, bold=True, size=9, color=(255, 255, 255))
                    # 深蓝背景
                    shading = table.rows[0].cells[j]._element.get_or_add_tcPr()
                    shd = shading.makeelement(qn('w:shd'), {
                        qn('w:fill'): '2F5496',
                        qn('w:val'): 'clear',
                    })
                    shading.append(shd)
                # 数据行
                for r, row_data in enumerate(rows):
                    for c, val in enumerate(row_data[:len(header)]):
                        set_cell_font(table.rows[r + 1].cells[c], val, size=9)
                doc.add_paragraph()  # 表后空行
            table_buffer = []
            in_table = False
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 分隔线
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): '999999',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 标题
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # 无序列表
        if re.match(r'^- ', line):
            text = re.sub(r'^- ', '', line)
            p = doc.add_paragraph(style='List Bullet')
            # 清除默认run
            p.clear()
            # 处理粗体
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # 块引用
        if line.startswith('> '):
            text = line[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            left = pBdr.makeelement(qn('w:left'), {
                qn('w:val'): 'single',
                qn('w:sz'): '18',
                qn('w:space'): '8',
                qn('w:color'): 'E74C3C',
            })
            pBdr.append(left)
            pPr.append(pBdr)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # 普通段落
        add_styled_paragraph(doc, line, size=11)
        i += 1

    doc.save(output_path)
    print(f"转换完成: {output_path}")
    print(f"文件大小: {os.path.getsize(output_path) / 1024:.1f} KB")


if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)) or ".")
    convert_md_to_docx(INPUT, OUTPUT)
